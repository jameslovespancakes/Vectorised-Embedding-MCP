"""DOCX parser via python-docx. Yields (text, None) per logical block."""

from __future__ import annotations

from pathlib import Path
from typing import Iterator

from docx import Document


def parse(path: Path) -> Iterator[tuple[str, None]]:
    doc = Document(str(path))
    buffer: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            buffer.append(text)
    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                buffer.append(" | ".join(cells))
    if buffer:
        yield "\n\n".join(buffer), None
